#!/usr/bin/env python3
"""
AngaWatch Competitive Analysis & Content Strategy Report Generator
Energy Mentors "Power the Community" Competition
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Page margins ───
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ─── Colour palette ───
DEEP_GREEN   = RGBColor(0x1A, 0x53, 0x36)   # #1A5336
MID_GREEN    = RGBColor(0x27, 0x7A, 0x48)   # #277A48
ACCENT_GOLD  = RGBColor(0xD4, 0xAF, 0x37)   # #D4AF37
LIGHT_GRAY   = RGBColor(0xF2, 0xF2, 0xF2)   # #F2F2F2
DARK_GRAY    = RGBColor(0x33, 0x33, 0x33)   # #333333
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)
LINK_BLUE    = RGBColor(0x00, 0x56, 0xD2)   # #0056D2
RED          = RGBColor(0xC0, 0x39, 0x2B)   # #C0392B

# ─── Helper: set paragraph shading ───
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

# ─── Helper: shade table cell ───
def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ─── Helper: set cell borders ───
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        props = kwargs.get(edge, {})
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), props.get('val', 'none'))
        if 'sz' in props:
            tag.set(qn('w:sz'), str(props['sz']))
        if 'color' in props:
            tag.set(qn('w:color'), props['color'])
        tcBorders.append(tag)
    tcPr.append(tcBorders)

# ─── Helper: bold run ───
def bold_run(para, text, size=None, color=None, italic=False):
    run = para.add_run(text)
    run.bold = True
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ─── Helper: normal run ───
def normal_run(para, text, size=None, color=None, italic=False, bold=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

# ─── Helper: heading ───
def add_heading(doc, text, level=1, color=None, size=None):
    if level == 1:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        shade_paragraph(p, '1A5336')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = WHITE
    elif level == 2:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = MID_GREEN
        # Underline via bottom border on paragraph
    elif level == 3:
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = DEEP_GREEN
    return p

# ─── Helper: body paragraph ───
def add_body(doc, text, size=11, indent=False):
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK_GRAY
    return p

# ─── Helper: bullet ───
def add_bullet(doc, text, sub=False, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if sub:
        p.paragraph_format.left_indent = Inches(0.5)
    if bold_part:
        run = p.add_run(bold_part)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
        run2 = p.add_run(text)
        run2.font.size = Pt(10.5)
        run2.font.color.rgb = DARK_GRAY
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
        run.font.color.rgb = DARK_GRAY
    return p

# ─── Helper: callout box ───
def add_callout(doc, title, body, bg='E8F5E9', title_color=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    shade_cell(cell, bg)
    cell.width = Inches(6.5)
    p1 = cell.add_paragraph()
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = title_color or DEEP_GREEN
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK_GRAY
    doc.add_paragraph()  # spacer

# ─── COVER PAGE ───────────────────────────────────────────────────────────────

cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(cover, '1A5336')
cover.paragraph_format.space_before = Pt(60)
cover.paragraph_format.space_after = Pt(4)
r = cover.add_run('ANGAWATCH (YATTA ECO-HUB)')
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = WHITE

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub1, '1A5336')
sub1.paragraph_format.space_before = Pt(0)
sub1.paragraph_format.space_after = Pt(4)
r2 = sub1.add_run('Competitive Analysis & Content Strategy')
r2.bold = True
r2.font.size = Pt(15)
r2.font.color.rgb = ACCENT_GOLD

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(sub2, '1A5336')
sub2.paragraph_format.space_before = Pt(0)
sub2.paragraph_format.space_after = Pt(60)
r3 = sub2.add_run('Energy Mentors · Power the Community Competition · 2025-2026')
r3.font.size = Pt(11)
r3.font.color.rgb = RGBColor(0xCC, 0xFF, 0xCC)

doc.add_paragraph()

# Meta info table
meta = doc.add_table(rows=3, cols=2)
meta.alignment = WD_TABLE_ALIGNMENT.CENTER
meta.style = 'Table Grid'
meta_data = [
    ('Prepared for', 'Nguitui Ryan — AngaWatch Team'),
    ('Competition Deadline', 'April 15, 2026'),
    ('Date', 'April 10, 2026'),
]
for i, (label, value) in enumerate(meta_data):
    row = meta.rows[i]
    shade_cell(row.cells[0], 'E8F5E9')
    p = row.cells[0].paragraphs[0]
    r = p.add_run(label)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = DEEP_GREEN
    p2 = row.cells[1].paragraphs[0]
    r2 = p2.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = DARK_GRAY

doc.add_page_break()

# ─── TABLE OF CONTENTS (manual) ───────────────────────────────────────────────

add_heading(doc, '  TABLE OF CONTENTS', level=1)
toc_items = [
    ('1', 'Competition Overview', '3'),
    ('2', 'Judging Framework Deep Dive', '4'),
    ('3', 'Energy Xcelerate — 2025 Grand Prize Winner', '5'),
    ('4', 'Past Winners Pattern Analysis', '7'),
    ('5', 'Competitive Gap Analysis: AngaWatch vs. Xcelerate', '8'),
    ('6', 'Strategic Recommendations for AngaWatch', '11'),
    ('7', 'Pitch Video Script (10–12 Minutes)', '13'),
    ('8', 'Slide Deck Strategy', '17'),
    ('9', 'Key URLs & Sources', '19'),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'{num}.  {title}')
    r.font.size = Pt(11)
    r.font.color.rgb = DARK_GRAY
    r.bold = (num == '1')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 1: COMPETITION OVERVIEW
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  1. COMPETITION OVERVIEW', level=1)

add_heading(doc, 'About Energy Mentors', level=2)
add_body(doc, 'Energy Mentors is a 501(c)(3) public charity founded by Don Victory (UConn 1981, retired ExxonMobil Chief Process Engineer with 40+ years of industry experience). Its mission is to catalyze the next generation of energy professionals through education, mentoring, and the "Power the Community" design competition. The organization is backed by founding sponsors EY, Aspen Technology (AspenTech), and Chart Industries, with MathWorks as software sponsor.')

add_heading(doc, 'Competition Concept', level=2)
add_body(doc, 'Power the Community is a global undergraduate and graduate student competition. Teams design energy infrastructure for a community of exactly 2,000 families whose income is defined as the combined salaries of a nurse and a schoolteacher in their chosen region — a purposeful affordability constraint that ensures designs are grounded in real-world economic reality rather than idealized high-budget scenarios.')

add_body(doc, 'Each submission must include a geolocation to an actual parcel of land, a community layout, a typical housing unit design, and a pitch video of 10–12 minutes (maximum 15 minutes) hosted on a publicly accessible platform such as YouTube or Vimeo. Deliverables must be in mass-market software formats (Microsoft, Google, Adobe).')

add_heading(doc, 'Prize Structure', level=2)

prize_tbl = doc.add_table(rows=4, cols=2)
prize_tbl.style = 'Table Grid'
prize_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
prize_headers = ['Award', 'Value']
for i, h in enumerate(prize_headers):
    cell = prize_tbl.rows[0].cells[i]
    shade_cell(cell, '1A5336')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)

prize_rows = [
    ('Grand Prize — Energy Innovators Award', '$10,000'),
    ('Distinguished Design Award (×2–3)', '$5,000–$7,500 each'),
    ('Total Prize Pool', '$31,000'),
]
for i, (award, val) in enumerate(prize_rows):
    row = prize_tbl.rows[i + 1]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'F2F2F2')
        shade_cell(row.cells[1], 'F2F2F2')
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(award)
    r0.font.size = Pt(10)
    r0.font.color.rgb = DARK_GRAY
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(val)
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = DEEP_GREEN

doc.add_paragraph()

add_heading(doc, 'Competition Timeline (2025–2026 Cycle)', level=2)
add_bullet(doc, 'Team Registration Deadline: March 15, 2026')
add_bullet(doc, 'Final Submission Deadline: April 15, 2026')
add_bullet(doc, 'Status: AngaWatch is within the active submission window')

add_heading(doc, 'Submission Requirements Checklist', level=2)
reqs = [
    'Geolocation of design to an actual parcel of land',
    'Statement of the most important needs of the community (with references)',
    'Benchmarks comparing the community vs. other communities/populations in the region',
    'Engineering, architectural, and economic calculations and renderings',
    'Design of a prototypical housing unit exemplifying the community\'s energy practices',
    'Community layout with focus on energy flows and infrastructure',
    'Pitch video: 10–12 minutes (max 15), publicly hosted (YouTube/Vimeo)',
    'All files in mass-market formats (MS Office, Google, Adobe)',
]
for r in reqs:
    add_bullet(doc, r)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 2: JUDGING FRAMEWORK
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  2. JUDGING FRAMEWORK DEEP DIVE', level=1)

add_heading(doc, 'No Formulistic Rubric — By Design', level=2)
add_body(doc, 'This is one of the most important things to understand about the competition: Energy Mentors deliberately chose NOT to use a point-based scoring rubric. The organizers believe that a formulistic rubric would bias submissions toward existing paradigms and constrain innovation. This is a critical strategic insight for AngaWatch.')

add_callout(doc,
    'Strategic Implication',
    'Because there is no fixed rubric, winning is about telling a compelling, coherent, and technically credible story — not checking boxes. Teams that communicate a clear vision, demonstrate deep local understanding, and show genuine innovation will stand out more than teams that optimize for a checklist.',
    bg='E8F5E9', title_color=DEEP_GREEN)

add_heading(doc, 'How Judging Actually Works', level=2)
add_bullet(doc, 'A panel of up to 5 independent industry judges reviews each group of 10–12 submissions')
add_bullet(doc, 'NEW judges are assigned for the finals (not the same as preliminary rounds)')
add_bullet(doc, 'Energy Mentors board members and employees cannot serve as judges')
add_bullet(doc, 'Known judges include: Jillian (Jill) Evanko (President & CEO, Chart Industries) and Vikas Dhole (GM, Sustainability, AspenTech)')
add_bullet(doc, 'A Lead Judge consolidates feedback from the panel')
add_bullet(doc, 'Judges identify: "Strong Elements" (cited by ≥2 judges), "Questions or Suggestions," and "Awesome" features')
add_bullet(doc, 'Grand Prize is decided by majority vote of judges')

add_heading(doc, 'What Judges Actually Look For (Inferred from Winners)', level=2)
add_body(doc, 'Analysing what past grand prize and distinguished design winners had in common reveals the de facto criteria that judges respond to:')

criteria_tbl = doc.add_table(rows=8, cols=2)
criteria_tbl.style = 'Table Grid'
criteria_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
c_header_cells = criteria_tbl.rows[0].cells
for i, h in enumerate(['What Judges Value', 'Evidence from Winners']):
    shade_cell(c_header_cells[i], '277A48')
    p = c_header_cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)

c_rows = [
    ('Deep local context knowledge', 'Winners consistently show intimate understanding of their specific community\'s challenges (grid unreliability, climate hazards, resource availability)'),
    ('Affordability within income constraint', 'The nurse + teacher salary constraint is non-negotiable; judges care about whether the design is economically feasible for the target household'),
    ('Innovation in energy mix', 'Winners used diverse, region-appropriate energy sources rather than generic solar-only designs'),
    ('Systems thinking beyond energy', 'Top designs addressed water, housing materials, internet access, food, and waste alongside energy'),
    ('Disaster/climate resilience', 'Multiple winners explicitly addressed regional climate hazards (hurricane-resistant homes, zero-outage year-round, drought resilience)'),
    ('Use of local/circular resources', 'Fly-ash concrete, EV battery repurposing, CBM from local sources — judges reward resource ingenuity'),
    ('Engineering & economic rigor', 'Credible calculations, not just concepts; specific IRR, cost per kWh, load analysis'),
]
for i, (criterion, evidence) in enumerate(c_rows):
    row = criteria_tbl.rows[i + 1]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'F9FBF9')
        shade_cell(row.cells[1], 'F9FBF9')
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(criterion)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = DEEP_GREEN
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(evidence)
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK_GRAY

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 3: ENERGY XCELERATE DEEP DIVE
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  3. ENERGY XCELERATE — 2025 GRAND PRIZE WINNER', level=1)

add_heading(doc, 'Team Profile', level=2)
add_bullet(doc, 'Team Name: Energy Xcelerate')
add_bullet(doc, 'Award: Energy Innovators Award — $10,000 Grand Prize (2025)')
add_bullet(doc, 'Members: Md Merajuddin Ahmed, Sanika Kole, Sachin Bangrawa, Abantika Duttagaur, Uddipan Mondal')
add_bullet(doc, 'Project Location: Lake Town, Kolkata, West Bengal, India')
add_bullet(doc, 'Scale: 2,000 families')

add_heading(doc, 'Project Design Overview', level=2)

add_heading(doc, 'Energy System', level=3)
e_tbl = doc.add_table(rows=4, cols=3)
e_tbl.style = 'Table Grid'
e_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, h in enumerate(['Energy Source', 'Share', 'Technology']):
    shade_cell(e_tbl.rows[0].cells[i], '1A5336')
    p = e_tbl.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(10)

e_data = [
    ('Solar', '67%', 'Rooftop + floating solar arrays with battery backup'),
    ('CBM (Compressed Biomethane / Coal Bed Methane)', '30%', 'Sourced from regional reserves; primary baseload'),
    ('Biogas', '3%', 'From community organic waste streams'),
]
for i, (src, pct, tech) in enumerate(e_data):
    row = e_tbl.rows[i + 1]
    if i % 2 == 0:
        for c in row.cells:
            shade_cell(c, 'F2F2F2')
    row.cells[0].paragraphs[0].add_run(src).font.size = Pt(10)
    p_pct = row.cells[1].paragraphs[0]
    r_pct = p_pct.add_run(pct)
    r_pct.bold = True
    r_pct.font.size = Pt(10)
    r_pct.font.color.rgb = DEEP_GREEN
    row.cells[2].paragraphs[0].add_run(tech).font.size = Pt(10)

doc.add_paragraph()

add_heading(doc, 'Housing & Community Design', level=3)
add_bullet(doc, 'Climate-responsive homes using mud blocks and passive cooling — appropriate for Kolkata\'s hot, humid climate and low-cost construction')
add_bullet(doc, 'Bioethanol-based cookstoves replacing LPG — cleaner indoor air quality, locally sourceable fuel')
add_bullet(doc, 'Rainwater harvesting and greywater reuse systems integrated into housing')
add_bullet(doc, 'Floating solar on water bodies — dual use of land/water, reduces evaporation')

add_heading(doc, 'Key Value Proposition', level=3)

add_callout(doc,
    '"Zero Outages Year-Round" — Xcelerate\'s Winning Message',
    'The team\'s central narrative was simple and powerful: a diversified renewable energy mix (solar + CBM + biogas) ensures that the community never loses power, regardless of season, weather, or grid instability. This addressed India\'s notorious power reliability problem head-on. Disaster resilience was the emotional core of their pitch.',
    bg='FFF8E1', title_color=RGBColor(0x7B, 0x5E, 0x00))

add_heading(doc, 'What Made Xcelerate Win — An Analysis', level=2)

win_factors = [
    ('1.', 'Clear, memorable headline', 'Every winning pitch needs a "so what" statement. Xcelerate\'s was "zero outages year-round." Judges remembered it.'),
    ('2.', 'Technically credible energy mix', 'The 67/30/3 split was precise, referenced, and matched to regional resources (CBM reserves in West Bengal, abundant solar radiation).'),
    ('3.', 'Local problem, local solution', 'Power outages are a real, painful, documented problem in Kolkata. The team showed they understood their community\'s life, not just its energy load.'),
    ('4.', 'Beyond energy: housing + water', 'Mud block passive cooling and greywater reuse showed judges that the team thought like community designers, not just engineers.'),
    ('5.', 'Disaster resilience narrative', 'India\'s monsoon season and cyclone exposure made disaster-proofing salient. Floating solar on water bodies was creative and regionally appropriate.'),
    ('6.', 'Strong team composition', 'Five members from engineering backgrounds brought multidisciplinary depth visible in both the energy system and housing design.'),
]
for num, factor, desc in win_factors:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True
    r1.font.size = Pt(10.5)
    r1.font.color.rgb = MID_GREEN
    r2 = p.add_run(f'{factor}: ')
    r2.bold = True
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = DARK_GRAY
    r3 = p.add_run(desc)
    r3.font.size = Pt(10.5)
    r3.font.color.rgb = DARK_GRAY

add_heading(doc, 'Pitch Video', level=2)
add_body(doc, 'A dedicated public pitch video for Energy Xcelerate was not located in open web searches at the time of this report. The competition requires teams to host their pitch on YouTube or Vimeo. If the video has been made public, it may be searchable on YouTube under "Energy Xcelerate Power the Community" or the team members\' names. The energymentors.org website also hosts winner highlights that may include embedded video links.')
add_bullet(doc, 'Search YouTube for: "Energy Xcelerate" OR "Md Merajuddin Ahmed" energy community')
add_bullet(doc, 'Check: https://www.energymentors.org/ptc-winners-2025')
add_bullet(doc, 'Check the Energy Mentors YouTube channel if one exists')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 4: PAST WINNERS PATTERN ANALYSIS
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  4. PAST WINNERS PATTERN ANALYSIS', level=1)

add_heading(doc, '2025 Competition Results (123 Teams, 320 Participants, 16+ Countries)', level=2)

p2025 = [
    ('Grand Prize', 'Energy Xcelerate', 'Lake Town, Kolkata, India',
     '$10,000', '67% solar + 30% CBM + 3% biogas; zero outages; mud block homes; floating solar'),
    ('Distinguished Design', 'AFTERLIFE', 'Rural Kenya',
     '$5,250', 'Solar-powered battery swapping system; repurposed EV cells; centralized charging for 2,000 homes — NOTE: also Kenya-based, direct competitive peer'),
    ('Distinguished Design', 'H2ENERGIES', 'Zvishavane, Zimbabwe',
     '$5,250', 'Hydrogen cookstoves + solar housing + Starlink internet + centralized water; holistic rural development'),
    ('Distinguished Design (reported)', 'Team Zambezi', 'Kariba, Zimbabwe',
     'N/A', 'Floating solar + biogas digesters + LiFePO4 batteries + AI energy management; hybrid community microgrid'),
]

tbl_2025 = doc.add_table(rows=len(p2025) + 1, cols=5)
tbl_2025.style = 'Table Grid'
for i, h in enumerate(['Award', 'Team', 'Location', 'Prize', 'Key Features']):
    shade_cell(tbl_2025.rows[0].cells[i], '277A48')
    p = tbl_2025.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(9)

for i, row_data in enumerate(p2025):
    row = tbl_2025.rows[i + 1]
    if i % 2 == 0:
        for c in row.cells:
            shade_cell(c, 'F9FBF9')
    for j, val in enumerate(row_data):
        p = row.cells[j].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(9)
        if j == 1:
            r.bold = True
        r.font.color.rgb = DARK_GRAY

doc.add_paragraph()

add_heading(doc, '2024 Competition Results (51 Teams, 107 Participants — Inaugural Year)', level=2)
add_bullet(doc, 'Grand Prize: "Creativity in Sustainability" — Shai Verma (India) & Rosemary Moyo (Zimbabwe), University of Connecticut', bold_part=None)
add_bullet(doc, 'Location: Mbizo, Kwekwe, Zimbabwe (Moyo\'s hometown)')
add_bullet(doc, 'Energy: Solar + batteries + natural gas cooking; concrete from fly ash (from nearby power plant)')
add_bullet(doc, 'Achievement: 27% CO2e community lifecycle reduction')
add_bullet(doc, 'Distinguished Design (Next Generation): Belrose Energy Company — Dominica; geothermal + wind; hurricane-resistant homes')
add_bullet(doc, 'Distinguished Design (Urban Planning): Jelly Fish — University of Houston Architecture students; Houston, TX')

add_heading(doc, 'Cross-Cycle Patterns: What Winners Have in Common', level=2)
patterns = [
    'Specific location: Winners always geolocate to a real, specific place — never generic. The specificity signals credibility.',
    'Named community need: The problem statement is personal and documented, not abstract.',
    'Resource-matched energy mix: Winners use what is actually available in their region (CBM in Bengal, geothermal in volcanic Dominica, fly ash near power plants in Zimbabwe).',
    'Housing material innovation: All top submissions designed housing that used local, climate-appropriate materials.',
    'Water and sanitation addressed: Grand prize winners consistently include water strategy.',
    'Circular economy thinking: Repurposed waste materials (fly ash, EV batteries, biogas from organic waste).',
    'Quantified impact: CO2 reduction, outage elimination, cost per kWh, affordability to target household.',
    '"Beyond energy" community vision: Schools, healthcare, agriculture, internet — top designs describe a community, not just a grid.',
]
for pat in patterns:
    add_bullet(doc, pat)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 5: COMPETITIVE GAP ANALYSIS
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  5. COMPETITIVE GAP ANALYSIS: ANGAWATCH vs. XCELERATE', level=1)

add_heading(doc, 'AngaWatch At-a-Glance', level=2)
aw_facts = [
    ('Location', 'Yatta Plateau, Machakos County, Kenya'),
    ('Land Area', '101.9 hectares (actual parcel)'),
    ('Scale', '2,000 homes'),
    ('Energy System', 'Grid-interactive microgrid — solar + biogas + BESS (Battery Energy Storage System)'),
    ('Target Household', 'Low-to-middle income Kenyan families (nurse + teacher income constraint satisfied)'),
    ('Capital Stack', 'KES 12,212 million'),
    ('Financial Return', '12–15% equity IRR'),
]
aw_tbl = doc.add_table(rows=len(aw_facts), cols=2)
aw_tbl.style = 'Table Grid'
for i, (k, v) in enumerate(aw_facts):
    row = aw_tbl.rows[i]
    if i % 2 == 0:
        shade_cell(row.cells[0], 'E8F5E9')
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.size = Pt(10)
    r0.font.color.rgb = DEEP_GREEN
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(v)
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK_GRAY

doc.add_paragraph()

add_heading(doc, 'Where AngaWatch is STRONGER than Xcelerate', level=2)

strengths = [
    (
        'More sophisticated financial modeling',
        'AngaWatch has a defined capital stack (KES 12,212M) and a projected equity IRR of 12–15%. Xcelerate\'s public description does not highlight equivalent financial depth. This is a major differentiator — it signals to industry judges that the team understands project finance, not just engineering.',
    ),
    (
        'Grid-interactive architecture (not off-grid)',
        'AngaWatch\'s grid-interactive microgrid is technically more advanced than a standalone off-grid design. It offers bidirectional grid connection, which enables resilience AND economic participation (potential energy export revenues). Xcelerate\'s design appears to be more self-contained.',
    ),
    (
        'BESS integration',
        'Battery Energy Storage System integration demonstrates awareness of grid stability, frequency regulation, and demand-side management — concepts that industry judges from AspenTech and Chart Industries will appreciate.',
    ),
    (
        'Real, specific Kenyan context',
        'Machakos County and the Yatta Plateau have specific, documented energy poverty challenges. Kenya\'s 2024 energy mix (primarily geothermal + hydro), rural electrification rates, and the Yatta area\'s semi-arid climate all provide a rich context that AngaWatch can leverage.',
    ),
    (
        'Kenya aligns with competition growth market',
        'The 2025 competition included Kenya-based teams (AFTERLIFE won a Distinguished Award). Judges are already familiar with and interested in Kenya. AngaWatch can build on this familiarity.',
    ),
    (
        'Larger, more detailed land plan',
        '101.9 ha is a specific, credible parcel. This level of detail in site selection signals genuine project development thinking.',
    ),
]
for title, desc in strengths:
    add_callout(doc,
        f'✅  {title}',
        desc,
        bg='E8F5E9', title_color=DEEP_GREEN)

add_heading(doc, 'Where AngaWatch is WEAKER or MISSING vs. Xcelerate', level=2)

gaps = [
    (
        'No single memorable headline / "zero outage" equivalent',
        'Xcelerate\'s "zero outages year-round" is a crisp, memorable claim that judges could repeat verbatim. AngaWatch needs an equally clear, bold, emotionally resonant headline. "Grid-interactive microgrid" is technically accurate but not emotionally compelling. What is the single promise AngaWatch makes to its 2,000 families?',
        'CRITICAL',
    ),
    (
        'Housing design not yet articulated',
        'Xcelerate designed mud block passive-cooling homes. Creativity in Sustainability used fly-ash concrete. AFTERLIFE repurposed EV batteries. AngaWatch must show a prototype housing unit that uses Kenyan vernacular architecture, local materials (e.g., stabilised compressed earth blocks — SCEBs, or stone from the Yatta region), and passive solar design for the semi-arid climate.',
        'CRITICAL',
    ),
    (
        'Water strategy gap',
        'The Yatta Plateau is semi-arid with rainfall of ~500–700mm/year and a history of water scarcity. Xcelerate included rainwater harvesting and greywater reuse. AngaWatch needs to explicitly address water — this is arguably the most acute community need in Machakos County and ignoring it could cost the top prize.',
        'HIGH',
    ),
    (
        'No clear disaster/climate resilience narrative',
        'Machakos County faces droughts, occasional flash floods, and land degradation. AngaWatch needs to frame how the microgrid and community design make residents resilient to these specific shocks.',
        'HIGH',
    ),
    (
        'Community needs statement not surfaced',
        'The submission requires an explicit "Statement of the most important needs of the community." For Yatta/Machakos: energy poverty (~30–40% rural electrification in the area), water scarcity, food insecurity (semi-arid agriculture), youth unemployment, and poor digital connectivity are all documented and citable. This statement must lead the submission.',
        'HIGH',
    ),
    (
        'Energy mix is not simplified / visualised',
        'The Xcelerate 67/30/3 split is visually memorable. AngaWatch\'s "solar + biogas + BESS" needs to be quantified: what % of load is met by solar? What is the biogas contribution? What is the BESS discharge duration? These numbers must be specific and memorable.',
        'MEDIUM',
    ),
    (
        '"Beyond energy" community vision is underdeveloped',
        'H2ENERGIES (Distinguished Award winner) included Starlink internet, centralized water, healthcare, and education. AngaWatch\'s Eco-Hub concept is promising but needs explicit development: What schools, clinics, markets, and agricultural facilities are in the 101.9 ha layout?',
        'MEDIUM',
    ),
    (
        'Benchmark data not yet cited',
        'Submissions require benchmarks comparing the community to other communities/populations in the region. AngaWatch should cite: Kenya rural electrification rate, Machakos County energy access data, average household energy spend, carbon emissions of kerosene/wood-fuel baseline, and how AngaWatch improves on all of these.',
        'MEDIUM',
    ),
    (
        'Differentiation from AFTERLIFE (Kenya competitor)',
        'AFTERLIFE (a 2025 Distinguished Design winner from Kenya) used solar battery swapping for rural homes. AngaWatch must clearly differentiate: grid-interactive microgrid vs. off-grid battery swapping, and why the AngaWatch model better serves the Yatta community.',
        'LOW-MEDIUM',
    ),
]
for title, desc, priority in gaps:
    bg = 'FDECEA' if priority == 'CRITICAL' else ('FFF3E0' if priority == 'HIGH' else 'F3F4F6')
    tc = RED if priority == 'CRITICAL' else (RGBColor(0xE6, 0x5C, 0x00) if priority == 'HIGH' else DARK_GRAY)
    add_callout(doc,
        f'⚠  [{priority}]  {title}',
        desc,
        bg=bg.lstrip('#') if not bg.startswith('#') else bg[1:], title_color=tc)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 6: STRATEGIC RECOMMENDATIONS
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  6. STRATEGIC RECOMMENDATIONS FOR ANGAWATCH', level=1)

add_heading(doc, 'Priority 1 — Craft a Headline Statement', level=2)
add_body(doc, 'Every winning submission has a headline that a judge can remember and repeat. Develop AngaWatch\'s single defining promise. Examples to workshop:')
add_bullet(doc, '"AngaWatch: Kenya\'s first grid-interactive eco-community — 2,000 families, zero energy bills, zero emissions."')
add_bullet(doc, '"The Yatta Eco-Hub: where 2,000 Kenyan families generate, store, and sell their own clean energy."')
add_bullet(doc, '"AngaWatch turns the Yatta Plateau from energy poverty to energy prosperity — sustainably, affordably, permanently."')
add_body(doc, 'Choose one and make it the first sentence of your video, your slide deck, and your written submission.')

add_heading(doc, 'Priority 2 — Design a Kenyan Vernacular Prototype Home', level=2)
add_body(doc, 'Design a single-family housing unit that:')
add_bullet(doc, 'Uses stabilised compressed earth blocks (SCEBs) or stone rubble masonry — both locally available in Machakos')
add_bullet(doc, 'Incorporates passive solar orientation (south-facing windows, roof overhangs for shading during hot season)')
add_bullet(doc, 'Integrates rooftop solar PV and a household BESS connection point')
add_bullet(doc, 'Includes a biogas connection for cooking and a water harvesting system')
add_bullet(doc, 'Targets a construction cost within reach of a nurse + teacher household budget in Kenya (~KES 2–4M for a modest 3-bedroom home)')

add_heading(doc, 'Priority 3 — Develop a Water & Sanitation Strategy', level=2)
add_body(doc, 'Machakos County has one of Kenya\'s most acute water scarcity challenges. AngaWatch must include:')
add_bullet(doc, 'Rainwater harvesting — Yatta average rainfall ~600mm/yr; calculate per-household tank sizing')
add_bullet(doc, 'Greywater recycling for irrigation of community food gardens')
add_bullet(doc, 'Solar-powered water pumping from the nearest reliable source (Athi River or borehole)')
add_bullet(doc, 'Water budget per household per day (WHO standard: 50L/person/day)')

add_heading(doc, 'Priority 4 — Quantify the Energy Mix', level=2)
add_body(doc, 'Replace "solar + biogas + BESS" with specific numbers. Recommended disclosure:')
add_bullet(doc, 'Solar PV: X MW peak, Y kWh/day generated, Z% of community demand')
add_bullet(doc, 'Biogas: from A tonnes/day of organic waste, B kW continuous baseload')
add_bullet(doc, 'BESS: C MWh storage, D hours of backup at full community load')
add_bullet(doc, 'Grid export potential: E kWh/day surplus sold back to Kenya Power grid')
add_bullet(doc, 'Average household energy bill: KES X/month (vs. KES Y/month current spend — a Z% saving)')

add_heading(doc, 'Priority 5 — Build the Community Layout Beyond Energy', level=2)
add_body(doc, 'Expand the 101.9 ha site plan to show:')
add_bullet(doc, 'Residential clusters (2,000 homes in neighbourhoods of ~50–100 homes each)')
add_bullet(doc, 'Community centre / Eco-Hub building (the "watch" hub — monitoring, data, community services)')
add_bullet(doc, 'Solar farm zone + biogas digester zone + BESS enclosure')
add_bullet(doc, 'Agri-zone: drip-irrigated food gardens using greywater')
add_bullet(doc, 'School, clinic, market / co-working space')
add_bullet(doc, 'Road network + EV charging (future-ready)')

add_heading(doc, 'Priority 6 — Prepare Benchmark Data', level=2)
add_body(doc, 'Cite these benchmarks in your submission (all are publicly available from Kenya Power, KNES, KNBS, and World Bank):')
add_bullet(doc, 'Kenya rural electrification rate: ~75% (2024), but Machakos rural areas lower')
add_bullet(doc, 'Average Kenyan household energy spend: KES 2,000–3,500/month (kerosene + charcoal + mobile charging)')
add_bullet(doc, 'Kenya grid emission factor: ~0.21 kgCO2e/kWh (primarily geothermal + hydro)')
add_bullet(doc, 'AngaWatch target: <0.05 kgCO2e/kWh — cite this reduction explicitly')
add_bullet(doc, 'Machakos County Human Development Index vs. Kenya average')

add_heading(doc, 'Priority 7 — Address Climate Resilience', level=2)
add_body(doc, 'Frame AngaWatch\'s resilience story around Machakos-specific hazards:')
add_bullet(doc, 'Drought: BESS + biogas as drought-proof baseload even when solar is reduced by dust or cloud')
add_bullet(doc, 'Flash floods: elevated BESS enclosures, permeable community roads, swales')
add_bullet(doc, 'Grid failure: microgrid islanding capability — community operates independently of Kenya Power during outages')
add_bullet(doc, 'Land degradation: Eco-Hub agri-zone restores soil health with biogas digestate as fertiliser')

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 7: PITCH VIDEO SCRIPT
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  7. PITCH VIDEO SCRIPT STRUCTURE (10–12 MINUTES)', level=1)

add_body(doc, 'The video is the primary judging deliverable. It must work as a standalone piece — a judge who reads nothing else must come away understanding the community\'s needs, your solution, and why it wins. The structure below is optimised for the competition\'s informal, innovation-focused judging style.')

add_callout(doc,
    'Video Platform Note',
    'Host on YouTube (unlisted or public). Ensure the link is stable and accessible without login. Title the video clearly: "AngaWatch | Yatta Eco-Hub | Energy Mentors Power the Community 2026". Include a description with team name, location, and submission date.',
    bg='E3F2FD', title_color=RGBColor(0x0D, 0x47, 0xA1))

scenes = [
    (
        'SCENE 1: THE HOOK',
        '0:00 – 1:30',
        '90 seconds',
        [
            'Open on aerial or drone footage of the Yatta Plateau — the actual 101.9 ha site if possible, or representative Machakos County landscape',
            'Voiceover: "The Yatta Plateau, Machakos County, Kenya. 2,000 families. A nurse. A schoolteacher. A community that works hard — but spends 30% of its income on unreliable, dirty energy."',
            'Cut to: A family at night using a kerosene lamp. A student struggling to study without reliable light. A health clinic losing power during a critical procedure.',
            'Voiceover: "What if this community could generate its own clean energy — and sell the surplus back to the grid?"',
            'Title card: AngaWatch. Yatta Eco-Hub. Machakos County, Kenya.',
            'KEY DATA: Show 3 stats: current household energy cost (KES/month), % of rural Machakos without reliable power, and AngaWatch\'s promise ("100% clean energy, 12 months/year")',
        ]
    ),
    (
        'SCENE 2: COMMUNITY NEEDS',
        '1:30 – 3:00',
        '90 seconds',
        [
            'Present the 5 core community needs, each with a brief visual and statistic:',
            '  — Reliable energy: Grid outages 8–12 hours/day in rural Machakos',
            '  — Affordable energy: Current spend KES 2,500+/month on fossil fuels and grid',
            '  — Clean water: Only X% of Yatta households have piped water',
            '  — Food security: Semi-arid climate, irregular rainfall, soil degradation',
            '  — Economic opportunity: Youth unemployment rate in Machakos County',
            'Benchmark slide: AngaWatch community vs. Machakos County vs. Kenya national average (table or bar chart)',
            'Close this scene with: "AngaWatch was designed to solve all five — simultaneously."',
        ]
    ),
    (
        'SCENE 3: THE SITE',
        '3:00 – 4:00',
        '60 seconds',
        [
            'Show the actual geolocation — map of Machakos County, zoom to Yatta Plateau, show the 101.9 ha parcel boundary',
            'Explain WHY this site: proximity to road network, solar irradiation data (GHI map of Kenya — Yatta scores well at ~5.5 kWh/m²/day), biogas feedstock from surrounding agriculture, Athi River water access',
            'Show the community site plan — a rendered top-down layout of the 101.9 ha with zones labeled',
            'KEY DATA: GPS coordinates, land area, solar resource data, distance to Kenya Power grid connection point',
        ]
    ),
    (
        'SCENE 4: ENERGY SYSTEM',
        '4:00 – 6:30',
        '150 seconds',
        [
            'Walk through the three-part energy system with animated or illustrated diagrams:',
            '  — Solar PV: X MW peak, rooftop on all 2,000 homes + community solar farm; show energy flow diagram',
            '  — Biogas: Community biogas digester fed by organic waste from 2,000 households + agri-zone; produces Y kW continuous power + cooking gas',
            '  — BESS: Z MWh battery storage; island mode capability; D hours of full-community backup',
            'Show the grid-interactive connection: AngaWatch sells surplus kWh back to Kenya Power; generates revenue that subsidises household energy bills',
            'Energy mix pie chart: Solar X%, Biogas Y%, Grid import Z% (target: <5%)',
            'KEY DATA: Average household energy bill post-AngaWatch (KES/month), CO2 reduction vs. baseline, BESS backup hours, grid export revenue estimate',
            'Close: "Unlike a standard off-grid system, AngaWatch is grid-interactive — meaning the community profits from its clean energy surplus."',
        ]
    ),
    (
        'SCENE 5: THE HOME — PROTOTYPE UNIT',
        '6:30 – 7:30',
        '60 seconds',
        [
            'Introduce the AngaWatch prototype home — rendered 3D exterior and floor plan',
            'Key features to highlight:',
            '  — Stabilised compressed earth blocks (SCEBs) — locally made from Machakos soil, low embodied carbon',
            '  — Passive solar orientation: south-facing windows, roof overhangs, thermal mass walls',
            '  — Integrated rooftop solar + household BESS connection point',
            '  — Biogas inlet for cooking + biogas lighting backup',
            '  — Rainwater harvesting cistern (roof catchment)',
            '  — Connection to community greywater recycling system',
            'KEY DATA: Home construction cost estimate (KES), monthly energy bill projection, CO2 saved vs. a conventional Kenyan rural home',
        ]
    ),
    (
        'SCENE 6: COMMUNITY LAYOUT & ECO-HUB',
        '7:30 – 8:30',
        '60 seconds',
        [
            'Walk through the 101.9 ha community plan — use an annotated site map or 3D render:',
            '  — Residential clusters: 20 × 100 homes, organised around shared green spaces',
            '  — Eco-Hub Centre: community monitoring & control room, co-working space, market, health post',
            '  — Solar farm + biogas digester zone (northeast, prevailing wind for odour management)',
            '  — BESS enclosure: elevated, flood-safe',
            '  — Agri-zone: drip-irrigated food gardens using greywater + biogas digestate fertiliser',
            '  — School + clinic',
            '  — Internal road network + EV charging points (future-ready)',
            'The "AngaWatch" monitoring system: real-time energy dashboard visible in Eco-Hub and on resident smartphones',
        ]
    ),
    (
        'SCENE 7: FINANCIALS',
        '8:30 – 9:45',
        '75 seconds',
        [
            'Present the capital stack clearly and confidently — this is a major differentiator for AngaWatch:',
            '  — Total capital: KES 12,212 million',
            '  — Equity: X% (KES Xm)',
            '  — Debt: Y% (KES Ym) — Green bonds / DFI financing (e.g., AIIB, IFC)',
            '  — Grants/concessional: Z% (KES Zm) — Kenya government / climate funds',
            'Equity IRR: 12–15% — show sensitivity table (base, upside, downside cases)',
            'Household affordability: Monthly energy fee = KES X (compare to current spend KES Y; saving = KES Z/month per family)',
            'Revenue streams: household energy subscriptions + grid export sales + carbon credits',
            'Payback period: X years',
            'Close: "AngaWatch is not just a community design — it is a bankable project."',
        ]
    ),
    (
        'SCENE 8: CLIMATE RESILIENCE & IMPACT',
        '9:45 – 10:45',
        '60 seconds',
        [
            'Address Machakos-specific climate challenges:',
            '  — Drought: BESS + biogas ensure 24/7 power even in zero-solar-days; drought-resilient agri-zone',
            '  — Flash floods: elevated infrastructure, permeable roads, community storm drainage',
            '  — Grid failure: Islanding capability — community runs independently when Kenya Power fails',
            'Impact metrics to display on screen:',
            '  — X tCO2e avoided per year',
            '  — Y households lifted from energy poverty',
            '  — Z litres of water saved/recycled per day',
            '  — KES X million in household energy savings per year across 2,000 families',
            '  — A jobs created (construction + operations)',
        ]
    ),
    (
        'SCENE 9: THE CALL TO ACTION / CLOSE',
        '10:45 – 12:00',
        '75 seconds',
        [
            'Return to the opening image of the Yatta Plateau — now with a rendered overlay of the AngaWatch community as it will look when built',
            'Team members on camera (optional but powerful): Each member briefly introduces themselves and their contribution',
            'Final statement: "AngaWatch is the Yatta Plateau\'s answer to energy poverty, climate risk, and economic exclusion — a community where 2,000 Kenyan families generate, store, and sell their own clean energy. This is not a concept. This is a plan."',
            'End card with: AngaWatch logo, team name, Yatta Plateau | Machakos County | Kenya, and the competition branding (Energy Mentors Power the Community 2026)',
            'Include submission details, contact, and the live YouTube link in the video description',
        ]
    ),
]

for scene_name, timing, duration, bullets in scenes:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    shade_paragraph(p, 'E8F5E9')
    r1 = p.add_run(f'{scene_name}   ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = DEEP_GREEN
    r2 = p.add_run(f'[{timing}  ·  {duration}]')
    r2.font.size = Pt(10)
    r2.font.color.rgb = MID_GREEN
    for b in bullets:
        add_bullet(doc, b, sub=b.startswith('  —'))

doc.add_paragraph()
add_callout(doc,
    'Production Tips for the Video',
    'Keep narration clear and confident. Use a mix of: (1) drone/location footage or Google Earth Pro zoom, (2) animated energy diagrams (PowerPoint/Canva animations exported to video), (3) rendered 3D site plan, (4) clear data slides with large fonts and minimal text, (5) optional: brief on-camera team introduction. You do NOT need professional production — clear audio and a logical structure matter more than camera quality. Judges have watched hundreds of presentations.',
    bg='E3F2FD', title_color=RGBColor(0x0D, 0x47, 0xA1))

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 8: SLIDE DECK STRATEGY
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  8. SLIDE DECK STRATEGY', level=1)

add_body(doc, 'The slide deck accompanies or supports the pitch video. It may be submitted as a separate deliverable. It should stand alone — a judge who only reads the deck should come away with the same understanding as one who watches the video. Aim for 25–35 slides total, structured as follows:')

slides = [
    ('Slides 1–2', 'Cover + Team', 'Title: "AngaWatch — Yatta Eco-Hub." Project tagline. Map showing Machakos County. Team photo (if available) + member names and backgrounds. One sentence each on their discipline.'),
    ('Slide 3', 'The Problem', 'Three statistics, large font, one per visual: (1) % of Yatta/Machakos without reliable power, (2) average household energy cost KES/month, (3) CO2 output of current energy mix. Single headline: "2,000 Kenyan families are one power cut away from crisis."'),
    ('Slide 4', 'Community Needs Statement', 'The 5 needs with icons: Energy / Water / Food / Income / Connectivity. One sentence per need with a supporting data point. Required by competition guidelines.'),
    ('Slide 5', 'Benchmarks Table', 'Table: Yatta community vs. Machakos County average vs. Kenya national average vs. WHO minimum standards, across 5 indicators. Required by competition guidelines.'),
    ('Slide 6', 'The Site', 'Map of Kenya → Machakos County → Yatta Plateau → 101.9 ha parcel. Show GPS coordinates. Note: actual parcel geolocation is a submission requirement.'),
    ('Slide 7', 'Site Advantages', 'Solar irradiation (GHI data), proximity to Athi River, road access, grid connection point, agricultural land quality. Use icons and concise callouts.'),
    ('Slides 8–9', 'Community Site Plan', 'Full-page rendered top-down site layout with zones labeled. Second slide: same map with energy flow arrows overlaid — solar to homes, biogas to digester, BESS to grid, surplus to Kenya Power.'),
    ('Slides 10–11', 'Energy System Diagram', 'Animated or static energy flow diagram showing all three sources + BESS + grid connection. Second slide: energy mix pie chart (X% solar, Y% biogas, Z% grid) with supporting load analysis.'),
    ('Slides 12–13', 'Solar Deep Dive', 'Total PV capacity (MW), annual generation (MWh/yr), panel technology, rooftop vs. ground-mount split. Solar irradiation map of the site. Show seasonal generation curve.'),
    ('Slides 14–15', 'Biogas Deep Dive', 'Feedstock: organic waste from 2,000 households + agri-zone. Digester capacity (m³), gas yield (m³/day), electricity generation (kW), digestate use as fertiliser. Show circular economy diagram.'),
    ('Slides 16–17', 'BESS + Grid-Interactive Architecture', 'BESS: capacity (MWh), discharge duration (hours at full load), chemistry (LFP recommended). Grid connection: islanding capability, export capacity (MW), annual export revenue (KES). THIS IS A KEY DIFFERENTIATOR — give it space.'),
    ('Slide 18', 'Prototype Home', '3D render exterior + floor plan. Callouts: SCEB walls, passive solar, rooftop PV, biogas inlet, rainwater cistern. Construction cost estimate. Monthly energy bill (KES X — vs. current KES Y).'),
    ('Slide 19', 'Water Strategy', 'Rainwater harvesting design (roof area × rainfall × collection efficiency = litres/day). Solar-powered borehole/pump. Greywater recycling for irrigation. Water budget per household (litres/day).'),
    ('Slide 20', 'Eco-Hub Centre', 'Render or sketch of the central Eco-Hub building. Features: energy monitoring & control room, co-working space, market, health post, connectivity hub. This is the "heart" of the community — name it, brand it.'),
    ('Slide 21', 'Climate Resilience', 'Three climate threats with icons: Drought / Flood / Grid Failure. For each: the threat, how AngaWatch mitigates it. Close: "AngaWatch is designed to thrive under Machakos climate conditions — not despite them."'),
    ('Slides 22–24', 'Financial Model', 'Slide 22: Capital stack (pie chart — equity / debt / grants; KES amounts). Slide 23: Revenue streams (household subscriptions + grid export + carbon credits); annual revenue (KES). Slide 24: Key metrics — Equity IRR 12–15%, Payback period X yrs, Household monthly fee KES X, NPV at 10% discount rate.'),
    ('Slide 25', 'Affordability Check', 'Show that the monthly energy fee is within reach of the nurse + teacher income target. Cite typical Kenyan nurse salary + schoolteacher salary. Show that the fee represents <10% of household income — competition\'s implicit test.'),
    ('Slide 26', 'Impact Summary', 'The "so what" slide. Five icons with numbers: tCO2e avoided / MWh clean energy generated / households electrified / water litres saved / jobs created. Make these big, visual, and memorable.'),
    ('Slides 27–28', 'Implementation Roadmap', 'Phased development: Phase 1 (Year 1–2: site prep, solar farm, first 500 homes), Phase 2 (Year 3: biogas + BESS, next 1,000 homes), Phase 3 (Year 4–5: full 2,000 homes, Eco-Hub complete). Gantt chart or timeline visual.'),
    ('Slide 29', 'Team', 'Team members: names, backgrounds, roles in the project. One line each. Photo if available.'),
    ('Slides 30–31', 'Appendix', 'Technical specifications, references, data sources, financial model assumptions. These are for judges who want to dig deeper.'),
]

slide_tbl = doc.add_table(rows=len(slides) + 1, cols=3)
slide_tbl.style = 'Table Grid'
slide_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, h in enumerate(['Slides', 'Section', 'Content & Visuals']):
    shade_cell(slide_tbl.rows[0].cells[i], '1A5336')
    p = slide_tbl.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = WHITE
    r.font.size = Pt(9)

for i, (nums, title, content) in enumerate(slides):
    row = slide_tbl.rows[i + 1]
    if i % 2 == 0:
        for c in row.cells:
            shade_cell(c, 'F9FBF9')
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(nums)
    r0.bold = True
    r0.font.size = Pt(9)
    r0.font.color.rgb = DEEP_GREEN
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = DARK_GRAY
    p2 = row.cells[2].paragraphs[0]
    r2 = p2.add_run(content)
    r2.font.size = Pt(9)
    r2.font.color.rgb = DARK_GRAY

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════
# SECTION 9: SOURCES
# ═══════════════════════════════════════════════════════════════════
add_heading(doc, '  9. KEY URLS & SOURCES', level=1)

sources = [
    ('Energy Mentors Official Website', 'https://www.energymentors.org/'),
    ('Competition Main Page', 'https://www.energymentors.org/power-the-community-competition'),
    ('Competition Basics & Submission Guidelines', 'https://www.energymentors.org/power-the-community-competition-basics'),
    ('2025 Winners Page', 'https://www.energymentors.org/ptc-winners-2025'),
    ('Participate / Register', 'https://www.energymentors.org/participate'),
    ('Judging Panel', 'https://www.energymentors.org/about-8'),
    ('UConn News — 2024 Winners (Creativity in Sustainability)', 'https://today.uconn.edu/2025/01/commitment-to-energy-sustainability-links-huskies-from-different-generations-through-international-contest/'),
    ('Opportunity Desk Competition Listing', 'https://opportunitydesk.org/2025/08/21/energy-mentors-power-the-community-international-college-design-competition-2025-2026-10000-prize/'),
    ('AspenTech Blog — Energy Mentors', 'https://www.aspentech.com/en/resources/blog/energy-mentors-shaping-next-generation-of-energy-leaders'),
]

for name, url in sources:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{name}: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = DARK_GRAY
    r2 = p.add_run(url)
    r2.font.size = Pt(10)
    r2.font.color.rgb = LINK_BLUE

doc.add_paragraph()
add_callout(doc,
    'Note on Energy Xcelerate Pitch Video',
    'A public pitch video for Energy Xcelerate was not located during research. To find it: (1) Search YouTube for "Energy Xcelerate Power the Community"; (2) Check the 2025 winners page for embedded links; (3) Search for team member names (Md Merajuddin Ahmed, Sanika Kole). If you find the video, analyse it for: opening hook, data density, visual production quality, and length of each section.',
    bg='FFF8E1', title_color=RGBColor(0x7B, 0x5E, 0x00))

# ─── Save ────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\phant\Desktop\Energy Mentors\AngaWatch_Competitive_Analysis.docx'
doc.save(output_path)
print(f'Saved: {output_path}')